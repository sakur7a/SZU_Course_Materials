from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "outputs"
FIG_DIR = OUT_DIR / "figures"
REPORT_DOCX = OUT_DIR / "王锦政_2024270223_算设实验3_地图填色报告.docx"
REPORT_MD = OUT_DIR / "实验3_地图填色问题实验报告.md"


def read_rows() -> list[dict[str, str]]:
    with (OUT_DIR / "summary.csv").open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 8.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    if widths is None:
        widths = [16.0 / len(headers)] * len(headers)
    for i, header in enumerate(headers):
        table.rows[0].cells[i].width = Cm(widths[i])
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=8.5)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Cm(widths[i])
            set_cell_text(cells[i], value, size=8)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        run.font.color.rgb = RGBColor(31, 78, 121)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def add_result_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    run.italic = True


def build_docx(rows: list[dict[str, str]]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("实验3：回溯法与地图填色问题")
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = subtitle.add_run("平面图四色定理、小规模验证、DIMACS 地图数据着色与随机平面图效率分析")
    srun.font.name = "宋体"
    srun._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    srun.font.size = Pt(10.5)

    add_heading(doc, "一、实验内容及过程", 1)
    add_para(doc, "本实验研究无向图的顶点着色问题。把地图中的每个区域抽象为图的一个顶点，若两个区域有公共边界，则在对应顶点之间连一条边。地图填色转化为：给定颜色数 k，为每个顶点分配一种颜色，使任意相邻顶点颜色不同。平面图满足四色定理，即任意平面图都可以用不超过 4 种颜色完成合法填色。")
    add_para(doc, "算法部分采用两类方法：小规模图和随机平面图使用 DSATUR 回溯。该方法每次优先选择饱和度最高的未着色顶点，并配合前向检查剪枝；450 顶点的 DIMACS 数据使用 TabuCol 局部搜索，从贪心初始解出发，在冲突顶点上移动颜色，并用禁忌表避免短周期震荡。")
    add_para(doc, "实验流程为：读取地图数据文件夹中的 .col 文件；构造粘贴图像对应的小规模邻接图；分别执行 4 色正确性测试、450 顶点实例指定颜色数测试，以及随机生成不同规模平面图后的 4 色效率测试；最后使用冲突边数验证染色是否合法。")

    add_heading(doc, "二、实验结果及分析", 1)
    add_heading(doc, "2.1 小规模地图正确性测试", 2)
    add_para(doc, "粘贴图像被抽象为 10 个区域、19 条邻接边的小图。DSATUR 回溯在一次搜索中完成 4 色填色，最终冲突边数为 0，说明建模和合法性检测逻辑正确。")
    doc.add_picture(str(FIG_DIR / "small_graph_coloring.png"), width=Cm(12.5))
    add_caption(doc, "图1 小规模粘贴地图抽象图的 4 色结果")

    small = rows[0]
    add_result_line(doc, f"小图：顶点数 {small['vertices']}，边数 {small['edges']}，颜色数 {small['colors']}，算法 {small['method']}，耗时 {small['elapsed']} s，冲突边 {small['conflicts']}。")

    add_heading(doc, "2.2 三个地图数据文件着色结果", 2)
    map_rows = [r for r in rows if r["name"].endswith(".col")]
    for r in map_rows:
        status = "成功" if r["success"] == "True" else "失败"
        add_result_line(doc, f"{r['name']}：顶点 {r['vertices']}，边 {r['edges']}，指定 {r['colors']} 色，{status}，耗时 {r['elapsed']} s，冲突边 {r['conflicts']}。")
    add_para(doc, "三个 450 顶点实例均在指定颜色数下找到合法染色，冲突边数均为 0。其中 le450_15b 边数最多且指定 15 色，搜索耗时相对更高；le450_25a 可用颜色更多，贪心初始解已经非常接近合法解，因此耗时最短。")
    doc.add_picture(str(FIG_DIR / "map_runtime.png"), width=Cm(12.5))
    add_caption(doc, "图2 三个 DIMACS 地图数据的着色耗时")

    add_heading(doc, "2.3 随机平面图规模效率分析", 2)
    planar_rows = [r for r in rows if r["name"].startswith("随机平面图")]
    for r in planar_rows:
        add_result_line(doc, f"n={r['vertices']}：边数 {r['edges']}，4 色成功，耗时 {r['elapsed']} s，回溯赋值次数 {r['assignments']}。")
    add_para(doc, "随机平面图采用增量三角剖分方式生成，因此始终保持平面性。测试规模从 20 增加到 300，所有样例都能用 4 种颜色完成合法着色。由于该生成方式得到的图结构较规则，DSATUR 的选择顺序很快找到可行解，实验中赋值次数基本等于顶点数，耗时随规模增加呈上升趋势但增长较平缓。")
    doc.add_picture(str(FIG_DIR / "planar_runtime.png"), width=Cm(12.5))
    add_caption(doc, "图3 随机平面图 4 色填色耗时随顶点数变化")
    doc.add_picture(str(FIG_DIR / "planar_edges.png"), width=Cm(12.5))
    add_caption(doc, "图4 随机平面图边数随顶点数变化")

    add_heading(doc, "三、算法时间复杂度分析", 1)
    add_para(doc, "设图 G=(V,E)，n=|V|，m=|E|，颜色数为 k，最大度数为 Δ。以下从精确搜索和局部搜索两个层面分析三种算法的复杂度。")

    add_heading(doc, "3.1 朴素暴力回溯", 2)
    add_para(doc, "暴力回溯按固定顺序逐顶点着色，每层至多 k 个分支。最坏情况下需遍历完整搜索树，包含 O(k^n) 个节点。每个节点处检查邻居颜色需 O(Δ)，因此最坏总时间复杂度为 O(k^n·Δ)。这是图着色 NP 完全性的直接体现：搜索空间对顶点数呈指数增长。当颜色数 k ≥ Δ+1 时，贪心即可保证线性着色；但当 k 接近色数 χ(G) 时，搜索树急剧膨胀。")

    add_heading(doc, "3.2 DSATUR 回溯", 2)
    add_para(doc, "DSATUR 的最坏复杂度与暴力法相同，仍为 O(k^n·Δ)。但其实际效率远优于暴力法，原因在于：（1）饱和度优先策略使冲突被尽早发现——当某顶点饱和度达到 k 时立即回溯，避免深入无效子树；（2）前向检查在每次赋值后从邻居域中删除已用颜色，若产生空域则剪枝。理想情况下（每步饱和度恰好为 k-1），搜索退化为线性回溯，赋值次数 O(n)。实验中 k=4 的平面图上，DSATUR 赋值次数基本等于 n，实际复杂度接近 O(n²)。")

    add_heading(doc, "3.3 TabuCol 局部搜索", 2)
    add_para(doc, "TabuCol 不构造搜索树，而是迭代改进一个完整着色方案。贪心初始化 O(m)，单步迭代中：找冲突顶点 O(n)，对最多 80 个候选顶点各尝试 k 种颜色（利用 adj_color 矩阵 O(1) 计算冲突变化量），执行移动并更新邻接计数 O(Δ)。单步总代价 O(n+k)。设总迭代次数为 T，则算法总代价为 O(T·(n+k))。T 取决于图结构和颜色数：k ≥ χ(G)+1 时通常在 O(n²) 步内收敛，k = χ(G) 时搜索空间变窄迭代次数增加。禁忌表防止短周期震荡，停滞 150000 步后重启，至多重启 80 次。")

    add_heading(doc, "3.4 三种算法复杂度对比", 2)
    add_table(doc,
        ["算法", "最坏复杂度", "每步代价", "实际特点"],
        [
            ["暴力回溯", "O(k^n·Δ)", "O(Δ)", "固定顺序，无图结构利用"],
            ["DSATUR 回溯", "O(k^n·Δ)", "O(n+Δ)", "饱和度选点，赋值数常≈n"],
            ["TabuCol", "O(T·(n+k))", "O(n+k)", "局部搜索，T取决于收敛速度"],
        ],
        widths=[2.5, 3.0, 2.5, 8.0],
    )
    add_para(doc, "暴力回溯和 DSATUR 最坏复杂度相同，但 DSATUR 通过动态选点大幅减少实际搜索量。TabuCol 放弃精确搜索的完备性，以多项式代价的迭代步骤快速逼近可行解，更适合大规模实例。")

    add_heading(doc, "四、实验结论", 1)
    add_para(doc, "地图填色问题可以自然抽象为图的顶点着色问题，关键约束是相邻区域不能同色。小规模地图实验验证了回溯搜索和合法性检测的正确性；随机平面图实验符合四色定理，所有测试图均可用 4 色完成合法填色。")
    add_para(doc, "对 450 顶点数据，单纯回溯搜索难以直接处理，因此引入 TabuCol 局部搜索更适合中等规模图着色。实验中 le450_5a、le450_15b、le450_25a 分别在 5、15、25 种颜色限制下全部得到 0 冲突解。")
    add_para(doc, "从效率上看，图规模增大通常会提高搜索时间；边数越多、颜色越少，约束越强，搜索难度越大。实际实现中，DSATUR 适合用于小规模或结构较好的平面图验证，TabuCol 更适合用于较大规模图的可行染色搜索。")

    doc.save(REPORT_DOCX)


def build_markdown(rows: list[dict[str, str]]) -> None:
    map_rows = [r for r in rows if r["name"].endswith(".col")]
    planar_rows = [r for r in rows if r["name"].startswith("随机平面图")]
    lines = [
        "# 实验3：回溯法与地图填色问题",
        "",
        "## 实验内容及过程",
        "",
        "将地图区域抽象为无向图顶点，将公共边界抽象为边，目标是在给定颜色数下为顶点着色，使相邻顶点颜色不同。小规模图和随机平面图采用 DSATUR 回溯，450 顶点数据采用 TabuCol 局部搜索。",
        "",
        "## 实验结果及分析",
        "",
        "### 小规模地图",
        "",
        f"- 顶点数：{rows[0]['vertices']}，边数：{rows[0]['edges']}，4 色成功，冲突边数：{rows[0]['conflicts']}。",
        "",
        "### 三个地图数据",
        "",
        "| 数据 | 顶点 | 边 | 颜色 | 成功 | 耗时/s | 冲突边 |",
        "|---|---:|---:|---:|---|---:|---:|",
    ]
    for r in map_rows:
        lines.append(f"| {r['name']} | {r['vertices']} | {r['edges']} | {r['colors']} | {r['success']} | {r['elapsed']} | {r['conflicts']} |")
    lines.extend([
        "",
        "### 随机平面图",
        "",
        "| 顶点 | 边 | 耗时/s | 回溯赋值次数 |",
        "|---:|---:|---:|---:|",
    ])
    for r in planar_rows:
        lines.append(f"| {r['vertices']} | {r['edges']} | {r['elapsed']} | {r['assignments']} |")
    lines.extend([
        "",
        "## 实验结论",
        "",
        "实验验证了地图填色到图着色的建模过程。小规模地图与随机平面图均能用 4 色完成合法填色，三个 450 顶点数据也分别在 5、15、25 色下得到 0 冲突解。随着顶点数和边数增加，搜索时间总体上升；颜色数越宽松，局部搜索越容易收敛。",
    ])
    REPORT_MD.write_text("\n".join(lines), encoding="utf-8-sig")


def main() -> None:
    rows = read_rows()
    build_docx(rows)
    build_markdown(rows)
    print(REPORT_DOCX)
    print(REPORT_MD)


if __name__ == "__main__":
    main()
